import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x10, 0x24, 0x3E)
    return h

def add_table_styled(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

# Title
title = doc.add_heading('PhD Research Plan \u2014 Physics-Informed Mamba Digital Twin for Mecanum-Wheeled Omnidirectional Robots', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    'Candidate: Vishvesh Koranne',
    'Guide: Prof. Shital S. Chiddarwar',
    'Institute: VNIT Nagpur, Department of Mechanical Engineering',
    'Target: PhD Thesis + IMECE 2026 paper + 2 journal papers',
    'Date: 2026-07-17'
]:
    meta.add_run(line + '\n').font.size = Pt(11)

doc.add_paragraph('\u2014' * 50)

# AIM
add_heading_styled(doc, 'AIM', 1)
doc.add_paragraph(
    'To develop a resilient, data-driven control architecture for 4WD Mecanum-Wheeled Omnidirectional Robots (FMWOR) '
    'for dynamic real-world conditions by closing the per-wheel friction identifiability gap '
    'through a physics-informed Mamba-SSM digital twin.'
)

# GOAL
add_heading_styled(doc, 'GOAL', 1)
doc.add_paragraph(
    'Design and deploy a Physics-Informed Mamba State-Space-Model Neural Network architecture '
    'for real-time FMWOR control, robust to nonlinear frictional interaction, unmodeled slipping, '
    'and mechanical wear of mecanum wheel rollers.'
)

# OBJECTIVES
add_heading_styled(doc, 'OBJECTIVES (4 Thesis Contributions)', 1)

obj_headers = ['Obj', 'Description', 'Success Metric']
obj_rows = [
    ['O1', 'High-fidelity sim-to-real pipeline \u2014 Julia 39-D LuGre-Adamov ODE \u2192 274 GB dataset \u2192 MuJoCo cross-validation \u2192 custom 4WD platform sim',
     '\u2022 5,670 trajectories, 24 profiles \u00d7 3 \u03bc \u00d7 \u03c7 sweep\n\u2022 MuJoCo state RMSE \u2264 5% vs Julia\n\u2022 Custom sim calibrated to hardware \u2264 10% parameter error'],
    ['O2', 'Mamba-SSM PINN architecture \u2014 selective SSM backbone + LuGre physics loss + per-wheel (\u03bc, \u03c7) heads for forward dynamics + inverse friction ID',
     '\u2022 \u2264 1 ms inference on Jetson Orin (FP16)\n\u2022 Forward dynamics RMSE \u2264 5% envelope\n\u2022 \u03bc RMSE \u2264 0.05, \u03c7 RMSE \u2264 0.001 (whitelisted)'],
    ['O4', 'Hardware validation \u2014 zero-shot (floor \u03bc) \u2192 one-shot (payload/CoM) \u2192 online ID (roller wear) on custom 4WD',
     '\u2022 Phase 1: pose RMSE \u2264 2\u00d7 sim (\u03bc \u2208 [0.3, 0.8])\n\u2022 Phase 2: pose RMSE \u2264 3\u00d7 sim (\u00b111% payload, \u00b120% CoM)\n\u2022 Phase 3: \u03bc(t) tracking R\u00b2 \u2265 0.8 vs wear progression'],
    ['O3', 'Certifiability-by-design framework \u2014 output-bounded PINN + reachability + runtime monitor for SIL 1\u20132 evidence',
     '\u2022 Lipschitz-constrained Mamba blocks (spectral norm \u2264 1)\n\u2022 Runtime monitor FP \u2264 1%\n\u2022 Reachable-set over-approx \u2264 2\u00d7 true envelope'],
]
add_table_styled(doc, obj_headers, obj_rows, col_widths=[0.4, 3.0, 3.0])

doc.add_paragraph()

# TASK HIERARCHY
add_heading_styled(doc, 'TASK HIERARCHY (Non-Linear Time, with Risk Forks)', 1)

# O1
add_heading_styled(doc, 'O1: SIM-TO-REAL PIPELINE', 2)

add_heading_styled(doc, 'T1.1 Julia ODE refinement & dataset generation (COMPLETE)', 3)
add_bullet(doc, '39-D stiff ODE (TRBDF2, reltol 1e-8, 2 kHz)')
add_bullet(doc, '24 excitation profiles, \u03bc \u2208 {0.3, 0.5, 0.8}, \u03c7 \u2208 [0.001, 0.01]')
add_bullet(doc, 'Output: Arrow + JLD2 (5345/5670 whitelisted)')

add_heading_styled(doc, 'T1.2 MuJoCo cross-validation (youBot) [PARALLEL with T2.1]', 3)
add_bullet(doc, 'Replicate geometry + LuGre in MJCF')
add_bullet(doc, 'Pass/Fail: State RMSE \u2264 5% AND Slip RMSE \u2264 8%')

add_heading_styled(doc, '\U0001f374 FORK: MuJoCo\u2013Julia mismatch >5% [Risk R2]', 4)
add_bullet(doc, 'F1.2a: Debug contact model (roller handoff, Hertz params)')
add_bullet(doc, 'F1.2b: If unresolved \u2192 Use Julia data ONLY for PINN training')
add_bullet(doc, 'F1.2c: Restrict MuJoCo to custom platform sys-ID only')

add_heading_styled(doc, 'Deliverable: Cross-validation report with pass/fail decision', 3)

add_heading_styled(doc, 'T1.3 Custom 4WD platform system identification [DEPENDS ON HARDWARE]', 3)
add_bullet(doc, 'MuJoCo model \u2192 hardware parameter fitting (mass, inertia, friction)')
add_bullet(doc, 'Target: \u2264 10% parameter error vs physical measurements')

add_heading_styled(doc, '\U0001f374 FORK: Custom platform hardware delays [Risk R3]', 4)
add_bullet(doc, 'F1.3a: Parallel sim work continues (Julia + MuJoCo)')
add_bullet(doc, 'F1.3b: Hardware-in-loop with recorded Julia data')
add_bullet(doc, 'F1.3c: Defer hardware-dependent tasks; advance O2/O4')

add_heading_styled(doc, 'Deliverable: Calibrated MuJoCo model for custom platform', 3)

# O2
doc.add_page_break()
add_heading_styled(doc, 'O2: MAMBA-PINN ARCHITECTURE [STARTS AFTER T1.1, PARALLEL WITH T1.2/T1.3]', 2)

add_heading_styled(doc, 'T2.1 Architecture design & implementation [CAN START IMMEDIATELY]', 3)
add_bullet(doc, 'Encoder: 2-layer BiGRU (hidden=128) \u2192 latent z (roller rates)')
add_bullet(doc, 'Backbone: Mamba-SSM (4 layers, d_model=256, d_state=16, expand=2)')
add_bullet(doc, 'Heads: Physics (\u03bc, \u03c7 \u00d7 4) + Dynamics (39-D \u1e35)')
add_bullet(doc, 'Export: ONNX \u2192 TensorRT for Jetson Orin')

add_heading_styled(doc, 'T2.2 Physics-informed training (3 stages) [NEEDS T1.2 PASS/FAIL DECISION]', 3)
add_bullet(doc, 'Stage 1 Grounding: \u03bb_phys=0, supervised on 5345 whitelisted')
add_bullet(doc, 'Stage 2 Physics: \u03bb_phys=1.0, physics loss on full 5670 + pseudo-labels')
add_bullet(doc, 'Stage 3 Identification: \u03bb_id=5.0, force-reconstruction + bounds')
add_bullet(doc, 'Loss: L_data + \u03bb_phys L_phys + \u03bb_id L_id + \u03bb_bounds L_bounds')
add_bullet(doc, 'Targets: dyn RMSE \u2264 5%, \u03bc RMSE \u2264 0.05, \u03c7 RMSE \u2264 0.001')

add_heading_styled(doc, 'T2.3 Inference optimization & profiling [NEEDS T2.2 MODEL]', 3)
add_bullet(doc, 'Target: \u2264 1 ms on Orin NX (FP16)')

add_heading_styled(doc, '\U0001f374 FORK: Mamba inference >1 ms [Risk R1]', 4)
add_bullet(doc, 'F2.3a: Profile early (TensorRT, FP16, batch=1)')
add_bullet(doc, 'F2.3b: Knowledge distillation \u2192 2-layer Mamba (d_model=128)')
add_bullet(doc, 'F2.3c: Distillation loss: L_KD = MSE(student, teacher) + \u03bb L_phys')
add_bullet(doc, 'F2.3c: Fallback target: ~0.35 ms, ~1.2M params')

add_heading_styled(doc, 'Deliverable: ONNX/TensorRT model + latency report', 3)

# O4
doc.add_page_break()
add_heading_styled(doc, 'O4: HARDWARE VALIDATION [STARTS AFTER T2.3 MODEL READY + T1.3 CALIBRATED]', 2)

add_heading_styled(doc, 'T4.1 Phase 1 \u2014 Floor \u03bc variation (zero-shot)', 3)
add_bullet(doc, 'Surfaces: epoxy (0.3), concrete (0.5), carpet (0.6), steel (0.3)')
add_bullet(doc, 'Trajectories: octagon, coupled_v\u03c9, spiral_iso \u00d7 3 reps')
add_bullet(doc, 'Metrics: pose RMSE (% traj length) \u2264 2\u00d7 sim; \u03bc RMSE \u2264 0.05')
add_bullet(doc, 'Control: MPC (horizon=20, 1 kHz) using twin + ASMC fallback')

add_heading_styled(doc, 'T4.2 Phase 2 \u2014 Payload/CoM shift (one-shot)', 3)
add_bullet(doc, 'Payload: +5.5% (1 kg), +11% (2 kg); CoM offset: \u00b113% h, \u00b120% l')
add_bullet(doc, 'Trajectories: ellipse, figure8 \u00d7 2 reps')
add_bullet(doc, 'Adaptation: 5 trajectories \u2192 fine-tune \u03bc heads')
add_bullet(doc, 'Metrics: pose RMSE \u2264 3\u00d7 sim; \u03bc drift \u2264 0.03 vs baseline')

add_heading_styled(doc, 'T4.3 Phase 3 \u2014 Roller wear (online ID)', 3)
add_bullet(doc, 'Wear stages: 0% \u2192 25% \u2192 50% \u2192 75% \u2192 100% (sandpaper abrasion)')
add_bullet(doc, 'Trajectories: spin_creep, multisine50 \u00d7 2 reps (spin-rich)')
add_bullet(doc, 'Primary: \u03bc(t) tracking R\u00b2 \u2265 0.8 vs measured progression')
add_bullet(doc, 'Secondary: \u03c7(t) correlation with \u03bc(t) (optional, wear mechanism)')
add_bullet(doc, 'Exploratory: Motor current FFT at 12\u00d7 wheel freq \u2192 wear index')
add_bullet(doc, 'Metrics: pose RMSE growth \u2264 2\u00d7 per stage')

add_heading_styled(doc, '\U0001f374 FORK: Wear progression too slow [Risk R4]', 4)
add_bullet(doc, 'F4.3a: Accelerated wear rig: motorized abrasion, 24/7')
add_bullet(doc, 'F4.3b: Simulated wear in MuJoCo (vary roller geometry, \u03bc)')
add_bullet(doc, 'F4.3c: Pre/post caliper validation (12 rollers/wheel)')

add_heading_styled(doc, 'Deliverable: Hardware validation report + sim-to-real gap characterization', 3)

# O3
doc.add_page_break()
add_heading_styled(doc, 'O3: CERTIFIABILITY FRAMEWORK [INFORMED BY O4 RESULTS]', 2)

add_heading_styled(doc, 'T3.1 Lipschitz-constrained Mamba blocks [USES O4 MEASURED GAPS]', 3)
add_bullet(doc, 'Spectral normalization on Mamba projections (spectral norm \u2264 1)')
add_bullet(doc, 'Compute Lipschitz constants per layer (power iteration)')
add_bullet(doc, 'Calibrate \u03ba in monitor using O4 Phase 1\u20132 disturbance statistics')
add_bullet(doc, 'Output bounding proof for 20-step horizon')

add_heading_styled(doc, 'T3.2 Reachability analysis [VALIDATED ON O4 TRAJECTORIES]', 3)
add_bullet(doc, 'Zonotope propagation through PINN')
add_bullet(doc, 'Target: reachable-set over-approximation \u2264 2\u00d7 true envelope')
add_bullet(doc, 'Validation: Monte Carlo on O4 hardware data vs analytical bounds')
add_bullet(doc, 'Adjust horizon/partitioning based on O4 chatter frequencies')

add_heading_styled(doc, 'T3.3 Runtime monitor design & validation [DEPLOYED ON O4 HARDWARE]', 3)
add_bullet(doc, 'Monitor: ||x_meas - x_pred|| > \u03ba\u00b7\u03c3 \u2192 safe stop trigger')
add_bullet(doc, '\u03ba calibrated from O4 Phase 1\u20132 residual distributions')
add_bullet(doc, 'False positive target: \u2264 1% (measured on O4 nominal runs)')
add_bullet(doc, 'Fault injection: bit-flip, NaN, OOD input \u2192 detection rate')

add_heading_styled(doc, '\U0001f374 FORK: Certifiability evidence insufficient [Risk R5]', 4)
add_bullet(doc, 'F3.3a: Scope claim: "enables SIL 1\u20132 evidence for learned component"')
add_bullet(doc, 'F3.3b: Explicitly exclude: full PL d, Category 3 hw, FSA')
add_bullet(doc, 'F3.3c: Produce: architecture diagram, bounds report, monitor FP rate')

add_heading_styled(doc, 'Deliverable: Certifiability evidence package (Lipschitz + reachability + monitor) validated on hardware', 3)

# TIMING INDEPENDENCE MAP
doc.add_page_break()
add_heading_styled(doc, 'TIMING INDEPENDENCE MAP', 1)

timing_headers = ['Task', 'Can Start', 'Blocks', 'Parallel With']
timing_rows = [
    ['T1.1', 'Day 0', 'T1.2, T2.1', '\u2014'],
    ['T1.2', 'T1.1 done', 'T2.2 (data decision)', 'T2.1'],
    ['T2.1', 'T1.1 done', 'T2.2', 'T1.2, T1.3'],
    ['T1.3', 'Hardware ready', 'T4.1 (calibrated sim)', 'T2.1, T2.2'],
    ['T2.2', 'T1.2 decision + T2.1', 'T2.3', 'T1.3'],
    ['T2.3', 'T2.2 done', 'T4.1 (model)', '\u2014'],
    ['T4.1', 'T2.3 + T1.3', 'T4.2, T3.1 (gap data)', '\u2014'],
    ['T4.2', 'T4.1', 'T4.3, T3.1/3.2', '\u2014'],
    ['T4.3', 'T4.2', 'T3.2/3.3 (validation)', '\u2014'],
    ['T3.1', 'T4.1 (initial gaps)', 'T3.2, T3.3', 'T4.2, T4.3'],
    ['T3.2', 'T3.1 + T4.2 data', 'T3.3', 'T4.3'],
    ['T3.3', 'T3.1 + T4.3 data', '\u2014', '\u2014'],
]
add_table_styled(doc, timing_headers, timing_rows, col_widths=[0.6, 1.2, 1.5, 1.5])

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Critical path: ')
run.bold = True
p.add_run('T1.1 \u2192 T1.2 \u2192 T2.2 \u2192 T2.3 \u2192 T4.1 \u2192 T4.2 \u2192 T4.3 \u2192 T3.3')

p = doc.add_paragraph()
run = p.add_run('Parallel tracks: ')
run.bold = True
p.add_run('T2.1 runs during T1.2/T1.3; T3.1 starts early with T4.1 data and refines through T4.3')

# RISK REGISTRY
doc.add_page_break()
add_heading_styled(doc, 'RISK REGISTRY (Mapped to Task Forks)', 1)

risk_headers = ['Risk ID', 'Risk', 'Likelihood', 'Impact', 'Primary Fork Location']
risk_rows = [
    ['R1', 'Mamba inference >1 ms on Orin', 'Medium', 'High', 'O2 \u2192 T2.3 \U0001f374'],
    ['R2', 'MuJoCo\u2013Julia mismatch >5%', 'Low', 'Medium', 'O1 \u2192 T1.2 \U0001f374'],
    ['R3', 'Custom platform hardware delays', 'Medium', 'High', 'O1 \u2192 T1.3 \U0001f374'],
    ['R4', 'Wear progression too slow', 'Medium', 'Medium', 'O4 \u2192 T4.3 \U0001f374'],
    ['R5', 'Certifiability evidence insufficient', 'Low', 'Medium', 'O3 \u2192 T3.3 \U0001f374'],
]
add_table_styled(doc, risk_headers, risk_rows, col_widths=[0.6, 2.0, 0.8, 0.8, 1.5])

# OUT-OF-SCOPE
add_heading_styled(doc, 'OUT-OF-SCOPE (Explicit Boundaries)', 1)

oos_headers = ['Item', 'Reason']
oos_rows = [
    ['Full PL d / SIL 2 certification', 'Requires Category 3 hardware, safety PLC, FSA \u2014 platform-level, not learned-component'],
    ['Hardware fault injection (IEC 61508-2)', 'Requires safety-rated test lab, fault injectors, months campaign \u2014 validates platform, not PINN architecture'],
    ['Multi-robot coordination', 'Single-platform focus'],
    ['End-to-end RL/IL policy', 'Model-based architecture (twin + MPC/SMC)'],
    ['Battery/energy as primary objective', 'Byproduct of slip reduction'],
]
add_table_styled(doc, oos_headers, oos_rows, col_widths=[2.0, 4.5])

# JUSTIFICATION: O4 BEFORE O3
doc.add_page_break()
add_heading_styled(doc, 'JUSTIFICATION: O4 BEFORE O3 ORDERING', 1)

just_headers = ['Argument', 'Detail']
just_rows = [
    ['Empirical grounding', 'Certifiability artifacts (Lipschitz bounds, reachability over-approx, monitor FP rate) need measured sim-to-real gaps from O4 to be calibrated. Theoretical bounds without hardware validation are weak evidence.'],
    ['Thesis claim validation', 'Claim: "physics-informed structure enables SIL 1\u20132 evidence." This is proven by showing the runtime monitor works on real hardware with real disturbances \u2014 not just in simulation.'],
    ['Design iteration', 'O4 Phase 1\u20132 reveals actual OOD behavior, chatter, sensor noise. These inform monitor threshold \u03ba, Lipschitz tightening, reachability horizon. Doing O3 first = designing blind.'],
    ['Risk reduction', 'If O4 reveals fundamental sim-to-real gaps (e.g., unmodeled dynamics), O3 framework adapts. Reverse order risks building a framework for a simulator that doesn\'t match reality.'],
]
add_table_styled(doc, just_headers, just_rows, col_widths=[1.5, 5.0])

# JUSTIFICATION: T2.1 BEFORE T1.3
add_heading_styled(doc, 'JUSTIFICATION: ARCHITECTURE (T2.1) BEFORE CUSTOM PLATFORM SYS-ID (T1.3)', 1)
doc.add_paragraph('\u2022 T2.1 depends only on Julia dataset (T1.1 complete) \u2014 independent of hardware')
doc.add_paragraph('\u2022 T1.3 needs physical robot (fabrication, integration) \u2014 schedule risk')
doc.add_paragraph('\u2022 Parallelizing T2.1 during T1.3 wait time keeps critical path moving')

# JUSTIFICATION: HARDWARE FAULT INJECTION OUT OF SCOPE
add_heading_styled(doc, 'JUSTIFICATION: HARDWARE FAULT INJECTION OUT OF SCOPE', 1)

fault_headers = ['Requirement', 'What It Needs', 'Why Out of Scope']
fault_rows = [
    ['IEC 61508-2 fault injection', 'Safety-rated hardware platform, programmable fault injectors (clock glitch, power glitch, EMFI, pin-level), fault campaign documentation', 'Requires dedicated safety test lab, certified equipment, months of campaign time \u2014 not a research contribution'],
    ['ISO 13849-1 Category 3 validation', 'Dual-channel architecture, cross-monitoring, diagnostic coverage \u226590% on hardware', 'Our contribution is the learned component\'s software evidence (Lipschitz, monitor); hardware architecture is a separate engineering effort'],
    ['Thesis scope', 'We deliver: software evidence artifacts enabling SIL 1\u20132 for the learned component', 'Full hardware fault injection validates the platform, not the PINN architecture'],
]
add_table_styled(doc, fault_headers, fault_rows, col_widths=[1.5, 2.5, 2.5])

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document version 2.0 \u2014 revised task hierarchy with O4 before O3, non-linear timing, risk forks, and explicit out-of-scope justifications.')
run.italic = True
run.font.size = Pt(10)

doc.save('PhD_Research_Plan_Mecanum_PINN_Mamba.docx')
print('DOCX saved successfully!')